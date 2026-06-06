"""DOCX handler — translate Word documents using python-docx.

Extracts paragraphs (including headings), table cells as individual fragments.
Rebuild preserves all formatting (bold, italic, font size, color, heading styles, etc.)
Does NOT add lang attributes — only preserves original format.
"""

import os
import re
import shutil

from .base import BaseHandler, TextFragment

_docx = None


def _get_docx():
    global _docx
    if _docx is None:
        import docx as _d
        _docx = _d
    return _docx


# ── Skip patterns ───────────────────────────────────────────────────
_CODE_PATTERNS = (
    re.compile(r'^\s*[{}\[\];]'),
    re.compile(r'^\s*(import |from |def |class |if |for |while )'),
    re.compile(r'^\s*(public |private |void |int |string )'),
    re.compile(r'^\s*(#include|using namespace)'),
    re.compile(r'^\s*<[^>]+>'),
    re.compile(r'^\s*(SELECT|INSERT|UPDATE|DELETE|CREATE)\s', re.I),
)


def _is_code_like(text):
    if not text or len(text.strip()) < 3:
        return False
    lines = text.strip().split('\n')
    if len(lines) > 2:
        code_lines = sum(1 for l in lines if any(p.search(l) for p in _CODE_PATTERNS))
        if code_lines / len(lines) > 0.4:
            return True
    return False


def _is_translatable(text):
    stripped = text.strip()
    if not stripped or len(stripped) < 2:
        return False
    if not any(c.isalpha() for c in stripped):
        return False
    if _is_code_like(stripped):
        return False
    return True


# ── Run formatting helpers ──────────────────────────────────────────

def _copy_run_format(source_run, target_run):
    """Copy all formatting from source run to target run (XML level)."""
    from docx.oxml.ns import qn

    # Get or create rPr (run properties) on target
    source_rPr = source_run._r.find(qn("w:rPr"))
    if source_rPr is not None:
        # Clone the rPr element
        target_rPr = source_rPr.__class__(source_rPr.tag)
        target_rPr.attrib = dict(source_rPr.attrib)
        for child in source_rPr:
            target_rPr.append(child.__class__(child.tag, child.attrib, child.text))
        # Replace target's rPr
        existing = target_run._r.find(qn("w:rPr"))
        if existing is not None:
            target_run._r.remove(existing)
        target_run._r.insert(0, target_rPr)


def _copy_para_format(source_para, target_para):
    """Copy paragraph formatting (style, alignment, spacing, etc.)."""
    from docx.oxml.ns import qn

    # Copy paragraph style
    if hasattr(source_para, 'style') and source_para.style:
        target_para.style = source_para.style

    # Copy paragraph format properties (alignment, indent, spacing)
    src_pPr = source_para._p.find(qn("w:pPr"))
    if src_pPr is not None:
        tgt_pPr = target_para._p.find(qn("w:pPr"))
        if tgt_pPr is None:
            tgt_pPr = src_pPr.__class__(qn("w:pPr"))
            target_para._p.insert(0, tgt_pPr)
        # Copy alignment, spacing, indent
        for tag in ("w:spacing", "w:ind", "w:jc", "w:keepLines", "w:keepNext"):
            elem = src_pPr.find(qn(tag))
            if elem is not None:
                existing = tgt_pPr.find(qn(tag))
                if existing is not None:
                    tgt_pPr.remove(existing)
                tgt_pPr.append(elem.__class__(elem.tag, dict(elem.attrib)))


# ── Handler ─────────────────────────────────────────────────────────

class DocxHandler(BaseHandler):
    """Handler for .docx Word documents."""

    @staticmethod
    def supported_extensions():
        return [".docx"]

    def extract(self, file_path, skip_tags=None, bilingual=True):
        """Extract all translatable content: paragraphs, headings, TOC entries, table cells."""
        docx = _get_docx()
        doc = docx.Document(file_path)
        fragments = []

        # 1. Extract ALL paragraphs (headings, normal text, TOC entries)
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not _is_translatable(text):
                continue
            style_name = para.style.name if para.style else "Normal"
            is_toc = style_name.startswith("TOC")
            fragments.append(TextFragment(
                text=text,
                meta={"type": "paragraph", "index": i,
                       "style": style_name, "is_toc": is_toc},
            ))

        # 2. Extract table cells
        for ti, table in enumerate(doc.tables):
            for ri, row in enumerate(table.rows):
                for ci, cell in enumerate(row.cells):
                    text = cell.text.strip()
                    if not _is_translatable(text):
                        continue
                    # Skip duplicate merged cells
                    if ci > 0 and cell._tc is row.cells[ci - 1]._tc:
                        continue
                    fragments.append(TextFragment(
                        text=text,
                        meta={"type": "cell", "table": ti, "row": ri, "col": ci},
                    ))

        return fragments

    def rebuild(self, file_path, fragments, translations, bilingual,
                target_lang="zh-CN", **kwargs):
        """Write translations back, preserving all formatting."""
        docx = _get_docx()
        output_path = self._output_path(file_path)
        shutil.copy2(file_path, output_path)
        doc = docx.Document(output_path)

        # Build lookup
        trans_map = {}
        for frag, trans in zip(fragments, translations):
            trans_map[self._meta_key(frag.meta)] = trans

        # 1. Paragraphs — process in reverse for safe insertion
        para_inserts = []
        for i, para in enumerate(doc.paragraphs):
            key = f"paragraph_{i}"
            if key not in trans_map:
                continue
            trans = trans_map[key]
            meta = fragments[[self._meta_key(f.meta) for f in fragments].index(key)].meta if key in trans_map else {}

            if meta.get("is_toc"):
                # TOC entry: update cached field text
                self._update_toc_entry(para, trans)
                if bilingual:
                    para_inserts.append((para, trans))
            elif bilingual:
                para_inserts.append((para, trans))
            else:
                self._replace_para_text(para, trans)

        for para, trans in reversed(para_inserts):
            self._insert_para_after(para, trans)

        # 2. Table cells
        for ti, table in enumerate(doc.tables):
            for ri, row in enumerate(table.rows):
                for ci, cell in enumerate(row.cells):
                    key = f"cell_{ti}_{ri}_{ci}"
                    if key not in trans_map:
                        continue
                    trans = trans_map[key]
                    if bilingual:
                        self._append_cell_translation(cell, trans)
                    else:
                        self._replace_cell_text(cell, trans)

        doc.save(output_path)
        return output_path

    # ── Helpers ─────────────────────────────────────────────────────

    @staticmethod
    def _meta_key(meta):
        if meta["type"] == "paragraph":
            return f"paragraph_{meta['index']}"
        return f"cell_{meta['table']}_{meta['row']}_{meta['col']}"

    @staticmethod
    def _output_path(file_path):
        base, ext = os.path.splitext(file_path)
        return f"{base}_translated{ext}"

    @staticmethod
    def _update_toc_entry(para, text):
        """Update TOC cached text in XML field codes.

        Word TOC entries use <w:fldChar> elements with nested <w:t> text.
        This updates the display text while preserving the field structure.
        """
        from docx.oxml.ns import qn

        # Find all w:t elements in the paragraph's XML
        # and replace their text (preserving field code structure)
        text_elements = para._p.findall(f".//{qn('w:t')}")
        if text_elements:
            # First w:t usually contains the visible TOC text
            text_elements[0].text = text
        else:
            # Fallback: use para.runs
            if para.runs:
                para.runs[0].text = text
                for run in para.runs[1:]:
                    run.text = ""

    @staticmethod
    def _replace_para_text(para, text):
        """Replace paragraph text, preserving ALL run formatting."""
        if not para.runs:
            para.text = text
            return

        # Write translation to first run (preserving its formatting)
        first_run = para.runs[0]
        first_run.text = text
        # Clear remaining runs
        for run in para.runs[1:]:
            run.text = ""

    @staticmethod
    def _insert_para_after(para, text):
        """Insert a new paragraph after the given one, copying style + formatting."""
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        from docx.text.paragraph import Paragraph

        new_p = OxmlElement("w:p")

        # 1. Copy paragraph properties (style, alignment, spacing)
        src_pPr = para._p.find(qn("w:pPr"))
        if src_pPr is not None:
            # Deep clone pPr
            import copy
            tgt_pPr = copy.deepcopy(src_pPr)
            new_p.insert(0, tgt_pPr)

        # 2. Create text run with formatting from original's first run
        run = OxmlElement("w:r")

        # Copy run formatting (bold, italic, font, size, color) from first run
        if para.runs:
            src_rPr = para.runs[0]._r.find(qn("w:rPr"))
            if src_rPr is not None:
                import copy
                tgt_rPr = copy.deepcopy(src_rPr)
                run.append(tgt_rPr)

        # Add text
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = text
        run.append(t)
        new_p.append(run)

        # 3. Insert after original paragraph
        para._p.addnext(new_p)

        return Paragraph(new_p, para._parent)

    @staticmethod
    def _append_cell_translation(cell, text):
        """Append translation to a table cell, preserving cell formatting."""
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        p = OxmlElement("w:p")
        run = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = text
        run.append(t)
        p.append(run)
        cell._tc.append(p)

    @staticmethod
    def _replace_cell_text(cell, text):
        """Replace cell content, preserving first run's formatting."""
        if not cell.paragraphs:
            return
        first_para = cell.paragraphs[0]
        if first_para.runs:
            first_para.runs[0].text = text
            for run in first_para.runs[1:]:
                run.text = ""
        else:
            first_para.text = text
        # Clear extra paragraphs
        for para in cell.paragraphs[1:]:
            for run in para.runs:
                run.text = ""
