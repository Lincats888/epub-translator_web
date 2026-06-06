#!/usr/bin/env python3
"""EPUB Translator Web Server — FastAPI + SSE

Run:
    uvicorn server.server:app --host 127.0.0.1 --port 8080
"""

import asyncio
import base64
import io
import json
import os
import shutil
import sys
import time
import uuid
import zipfile
from concurrent.futures import ThreadPoolExecutor
from pathlib import Path
from threading import Lock
from xml.etree import ElementTree as ET

import yaml
from bs4 import BeautifulSoup
from fastapi import FastAPI, UploadFile, File, HTTPException, Query
from fastapi.responses import FileResponse, StreamingResponse, HTMLResponse, Response

# ── Import existing EpubTranslator modules (no modifications) ──────────
PROJECT_ROOT = str(Path(__file__).resolve().parent.parent)
if PROJECT_ROOT not in sys.path:
    sys.path.insert(0, PROJECT_ROOT)

from epub_translator.config import Config
from epub_translator.extractor import EpubExtractor
from epub_translator.cache import TranslationCache
from epub_translator.parser import parse_file
from epub_translator.translator import Translator, StopTranslation
from epub_translator.rebuilder import inject_line_height, rebuild_epub
from epub_translator.crypto import encrypt, decrypt, is_encrypted

from handlers import get_handler, get_supported_extensions, is_supported
from handlers.pdf_handler import PdfHandler
from languages import get_all_languages, get_lang_name, detect_language, get_system_prompt


def _doc_page_count(path):
    """Helper: count pages in a PDF."""
    try:
        import fitz
        doc = fitz.open(path)
        n = len(doc)
        doc.close()
        return n
    except Exception:
        return 40  # fallback

# ── App & State ────────────────────────────────────────────────────────
app = FastAPI(title="EPUB Translator")

_executor = ThreadPoolExecutor(max_workers=2)
_tasks: dict[str, dict] = {}
_lock = Lock()

TEMP_DIR = os.path.join(PROJECT_ROOT, "temp")
OUTPUT_DIR = os.path.join(PROJECT_ROOT, "output")
CONFIG_PATH = os.path.join(PROJECT_ROOT, "config.yaml")


def _update(task_id: str, **kwargs):
    with _lock:
        if task_id in _tasks:
            _tasks[task_id].update(kwargs)


# ── EPUB Metadata Reader ──────────────────────────────────────────────

NS = {
    "container": "urn:oasis:names:tc:opendocument:xmlns:container",
    "opf": "http://www.idpf.org/2007/opf",
    "dc": "http://purl.org/dc/elements/1.1/",
    "ncx": "http://www.daisy.org/z3986/2005/ncx/",
    "xhtml": "http://www.w3.org/1999/xhtml",
}


def _read_epub_meta(epub_path: str) -> dict:
    """Read EPUB metadata: title, author, cover image, TOC."""
    result = {"title": "", "author": "", "cover": None, "toc": [], "opf_dir": "OEBPS/"}

    with zipfile.ZipFile(epub_path, "r") as zf:
        # 1. Find OPF path from container.xml
        try:
            container = zf.read("META-INF/container.xml")
            root = ET.fromstring(container)
            opf_path = root.find(".//container:rootfile", NS).get("full-path")
        except Exception:
            opf_path = "OEBPS/content.opf"

        opf_dir = os.path.dirname(opf_path)
        result["opf_dir"] = opf_dir + "/" if opf_dir else ""

        # Helper: resolve manifest href to zip path (always forward slashes)
        def _zip_path(href):
            if opf_dir:
                return opf_dir + "/" + href
            return href

        # 2. Parse OPF
        try:
            opf = ET.fromstring(zf.read(opf_path))
        except Exception:
            return result

        # Title & Author
        title_el = opf.find(".//dc:title", NS)
        if title_el is not None and title_el.text:
            result["title"] = title_el.text.strip()
        author_el = opf.find(".//dc:creator", NS)
        if author_el is not None and author_el.text:
            result["author"] = author_el.text.strip()

        # Build manifest dict
        manifest = {}
        for item in opf.findall(".//opf:manifest/opf:item", NS):
            manifest[item.get("id", "")] = item

        # 3. Cover image — try multiple strategies
        cover_item = None

        # Strategy A: EPUB2 <meta name="cover" content="item-id"/>
        for meta in opf.findall(".//opf:meta", NS):
            if meta.get("name") == "cover":
                cid = meta.get("content", "")
                if cid in manifest and "image" in manifest[cid].get("media-type", ""):
                    cover_item = manifest[cid]
                break

        # Strategy B: EPUB3 properties="cover-image"
        if not cover_item:
            for item in manifest.values():
                if "cover-image" in item.get("properties", ""):
                    cover_item = item
                    break

        # Strategy C: filename contains "cover"
        if not cover_item:
            for item in manifest.values():
                href = item.get("href", "")
                if "image" in item.get("media-type", "") and "cover" in href.lower():
                    cover_item = item
                    break

        # Strategy D: first image in manifest
        if not cover_item:
            for item in manifest.values():
                if "image" in item.get("media-type", ""):
                    cover_item = item
                    break

        if cover_item is not None:
            href = cover_item.get("href", "")
            if href:
                try:
                    img_data = zf.read(_zip_path(href))
                    mime = cover_item.get("media-type", "image/jpeg")
                    b64 = base64.b64encode(img_data).decode()
                    result["cover"] = f"data:{mime};base64,{b64}"
                except KeyError:
                    pass

        # 4. TOC — NCX (EPUB2)
        ncx_href = None
        for item in manifest.values():
            if item.get("media-type") == "application/x-dtbncx+xml":
                ncx_href = item.get("href")
                break

        if ncx_href:
            try:
                ncx = ET.fromstring(zf.read(_zip_path(ncx_href)))
                def _parse_ncx(navpoint, depth=0):
                    items = []
                    label = navpoint.find(".//ncx:text", NS)
                    content = navpoint.find("ncx:content", NS)
                    href = content.get("src", "") if content is not None else ""
                    # Strip anchor fragments
                    href = href.split("#")[0]
                    title = label.text.strip() if label is not None and label.text else ""
                    items.append({
                        "title": title,
                        "href": href,
                        "depth": depth,
                    })
                    for child in navpoint.findall("ncx:navPoint", NS):
                        items.extend(_parse_ncx(child, depth + 1))
                    return items
                nav_map = ncx.find(".//ncx:navMap", NS)
                if nav_map is not None:
                    for nav in nav_map.findall("ncx:navPoint", NS):
                        result["toc"].extend(_parse_ncx(nav, 0))
            except Exception:
                pass

        # 5. TOC — nav.xhtml (EPUB3)
        def _parse_nav_ol(ol, depth=0):
            items = []
            for li in ol.find_all("li", recursive=False):
                a = li.find("a", href=True)
                title = a.get_text(strip=True) if a is not None else ""
                href = a["href"].split("#")[0] if a is not None and a.get("href") else ""
                items.append({"title": title, "href": href, "depth": depth})
                child_ol = li.find("ol")
                if child_ol:
                    items.extend(_parse_nav_ol(child_ol, depth + 1))
            return items

        if not result["toc"]:
            for item in manifest.values():
                if "nav" in item.get("properties", ""):
                    nav_href = item.get("href")
                    if nav_href:
                        try:
                            nav_html = zf.read(_zip_path(nav_href)).decode("utf-8", errors="replace")
                            soup = BeautifulSoup(nav_html, "lxml")
                            nav_el = soup.find("nav")
                            if nav_el:
                                ol = nav_el.find("ol")
                                if ol:
                                    result["toc"] = _parse_nav_ol(ol, 0)
                        except Exception:
                            pass
                    break

        # 6. TOC — fallback to spine order
        if not result["toc"]:
            for ref in opf.findall(".//opf:spine/opf:itemref", NS):
                href = manifest.get(ref.get("idref", ""), ET.Element("x")).get("href", "")
                if href and href.endswith((".html", ".xhtml", ".htm")):
                    # Use filename without extension as fallback title
                    title = os.path.splitext(os.path.basename(href))[0]
                    result["toc"].append({"title": title, "href": href, "depth": 0})

    return result


def _read_epub_content(epub_path: str, opf_dir: str, href: str) -> str:
    """Read a chapter's HTML content from EPUB, inlining images as base64 data URIs."""
    full_path = os.path.join(opf_dir, href).replace("\\", "/")
    with zipfile.ZipFile(epub_path, "r") as zf:
        try:
            html = zf.read(full_path).decode("utf-8", errors="replace")
        except KeyError:
            return "<p>Content not found.</p>"

        # Extract body content
        soup = BeautifulSoup(html, "lxml")

        # Strip links so they don't cause navigation errors in preview
        for a in soup.find_all("a", href=True):
            del a["href"]

        # Inline images as data URIs so they display in the preview
        chapter_dir = os.path.dirname(href).replace("\\", "/")
        for img in soup.find_all("img"):
            src = img.get("src", "")
            if not src or src.startswith(("http://", "https://", "data:")):
                continue
            # Resolve relative path against chapter location (normalize ..)
            if src.startswith("/"):
                img_rel_path = src.lstrip("/")
            else:
                img_rel_path = os.path.normpath(os.path.join(chapter_dir, src)).replace("\\", "/")
            img_zip_path = os.path.join(opf_dir, img_rel_path).replace("\\", "/")
            try:
                img_data = zf.read(img_zip_path)
                ext = os.path.splitext(src)[1].lower()
                mime_map = {".png": "image/png", ".jpg": "image/jpeg", ".jpeg": "image/jpeg",
                            ".gif": "image/gif", ".svg": "image/svg+xml", ".webp": "image/webp"}
                mime = mime_map.get(ext, "image/png")
                b64 = base64.b64encode(img_data).decode()
                img["src"] = f"data:{mime};base64,{b64}"
            except (KeyError, FileNotFoundError):
                pass  # Leave broken link as-is

        body = soup.find("body")
        if body:
            return str(body)
        return html


# ── Translation Runner ────────────────────────────────────────────────

def _run_translate(task_id: str, epub_path: str, target_lang: str = "zh-CN", bilingual: bool = True):
    """Background EPUB translation — mirrors main.py cmd_translate logic."""
    try:
        _update(task_id, status="loading", step="Loading configuration...")
        config = Config(CONFIG_PATH)
        config.load()
        if not config.api_key or config.api_key in ("sk-xxxx", "sk-your-api-key-here"):
            _update(task_id, status="error", error="API key not configured. Click settings.")
            return

        _update(task_id, step="Extracting EPUB...")
        extractor = EpubExtractor(epub_path, TEMP_DIR)
        extract_dir = extractor.extract()
        book_name = os.path.splitext(os.path.basename(epub_path))[0]

        cache = TranslationCache(extract_dir)
        cache.load()

        content_files = extractor.list_content_files()
        toc_file = extractor.find_toc_file()
        if toc_file and toc_file not in content_files:
            content_files.append(toc_file)
        opf_file = extractor.get_opf_path()
        if opf_file not in content_files:
            content_files.append(opf_file)

        if not content_files:
            _update(task_id, status="error", error="No content files found in EPUB")
            return

        total_files = len(content_files)
        translator_obj = Translator(config, target_lang, bilingual)
        total_translated = 0
        total_cached = 0
        is_bilingual = bilingual

        _update(task_id, status="translating", step="Translating...",
                current_file=0, total_files=total_files,
                file_progress=0, file_total=0)

        for file_idx, file_path in enumerate(content_files, 1):
            # Check stop flag
            if _tasks.get(task_id, {}).get("stopped"):
                _update(task_id, status="stopped", step="Stopped by user")
                return

            rel_path = os.path.relpath(file_path, extract_dir)
            parsed = parse_file(file_path, config.skip_tags, bilingual=is_bilingual)
            if not parsed.fragments:
                continue

            uncached_indices = []
            uncached_texts = []
            for i, frag in enumerate(parsed.fragments):
                cached = cache.get(frag.text)
                if cached is not None:
                    total_cached += 1
                else:
                    uncached_texts.append(frag.text)
                    uncached_indices.append(i)

            if uncached_texts:
                _update(task_id, current_file=file_idx, total_files=total_files,
                        file_name=rel_path, file_progress=0,
                        file_total=len(uncached_texts))

                def on_progress(done, _ft=len(uncached_texts), _fi=file_idx,
                                _tf=total_files, _rn=rel_path):
                    _update(task_id, current_file=_fi, total_files=_tf,
                            file_name=_rn, file_progress=min(done, _ft),
                            file_total=_ft)

                try:
                    translations_result = translator_obj.translate_all(
                        uncached_texts, progress_callback=on_progress,
                        stop_check=lambda: _tasks.get(task_id, {}).get("stopped", False)
                    )
                except StopTranslation:
                    cache.flush()
                    _update(task_id, status="stopped", step="Stopped by user")
                    return

                file_translations = [None] * len(parsed.fragments)
                for i, frag in enumerate(parsed.fragments):
                    cached = cache.get(frag.text)
                    if cached is not None:
                        file_translations[i] = cached
                for idx, translation in zip(uncached_indices, translations_result):
                    file_translations[idx] = translation
                    cache.put(parsed.fragments[idx].text, translation)
                    total_translated += 1

                cache.flush()
            else:
                file_translations = [cache.get(frag.text) for frag in parsed.fragments]

            parsed.save(file_translations)

        _update(task_id, step="Rebuilding EPUB...")
        if is_bilingual:
            inject_line_height(extract_dir)
        output_path = rebuild_epub(extract_dir, OUTPUT_DIR, book_name)

        _update(task_id, status="done", step="Complete!",
                output=output_path, translated=total_translated, cached=total_cached)

    except Exception as e:
        _update(task_id, status="error", error=str(e))


def _run_translate_ocr(task_id: str, file_path: str, target_lang: str = "zh-CN",
                        bilingual: bool = True):
    """OCR pipeline: Vision API OCR per page → translate → clean text pages."""
    try:
        _update(task_id, status="loading", step="Loading configuration...")
        config = Config(CONFIG_PATH)
        config.load()
        if not config.api_key or config.api_key in ("sk-xxxx", "sk-your-api-key-here"):
            _update(task_id, status="error", error="API key not configured. Click settings.")
            return

        from handlers.pdf_ocr import PdfOcr
        from handlers.pdf_scanned_handler import PdfScannedRebuilder

        import fitz as _fitz

        # ── Phase 1: Extract text per page ────────────────────────
        # Check if scanned → use Vision API OCR; if text PDF → direct extract
        _is_scanned = PdfOcr.is_scanned(file_path)
        _update(task_id, status="translating",
                step="Extracting text..." if not _is_scanned else "OCR scanning...",
                file_progress=0, file_total=100)

        if _is_scanned:
            ocr_api_key = config.ocr_api_key
            if not ocr_api_key:
                _update(task_id, status="error",
                        error="OCR API key not configured. Please set it in Settings.")
                return

            ocr = PdfOcr(
                api_key=ocr_api_key,
                base_url=config.ocr_api_base or "https://api.siliconflow.com/v1",
                model=config.ocr_model or "Qwen/Qwen3-VL-32B-Instruct",
            )

            page_texts = ocr.ocr_pdf(file_path,
                                     progress_callback=lambda d, t, txt=None: _update(
                                         task_id,
                                         file_progress=d, file_total=t,
                                         step=f"OCR {d}/{t}"))
        else:
            # Text PDF — extract directly with PyMuPDF (instant)
            _doc = _fitz.open(file_path)
            total_p = _doc.page_count
            page_texts = []
            for i in range(total_p):
                page_texts.append(_doc[i].get_text())
                _update(task_id,
                        file_progress=i + 1, file_total=total_p,
                        step=f"Extracting {i + 1}/{total_p}")
            _doc.close()

        total_pages = len(page_texts)

        # Filter empty pages
        non_empty = [t for t in page_texts if t.strip()]
        if not non_empty:
            _update(task_id, status="error",
                    error="OCR found no text on any page.")
            return

        # ── Phase 2: Translate ─────────────────────────────────────
        _update(task_id, status="translating",
                step=f"Translating {len(non_empty)} pages...",
                file_progress=total_pages, file_total=total_pages * 2)

        translator_obj = Translator(config, target_lang, bilingual)

        def on_trans_progress(done):
            _update(task_id,
                    file_progress=total_pages + done,
                    file_total=total_pages * 2,
                    step=f"Translating {total_pages + done}/{total_pages * 2}")

        page_translations = translator_obj.translate_all(
            non_empty, progress_callback=on_trans_progress,
            stop_check=lambda: _tasks.get(task_id, {}).get("stopped", False)
        )

        if _tasks.get(task_id, {}).get("stopped"):
            _update(task_id, status="stopped", step="Stopped by user")
            return

        # Map translations back to original page indices
        trans_map = {}
        ti = 0
        for pi in range(total_pages):
            if page_texts[pi].strip():
                trans_map[pi] = page_translations[ti]
                ti += 1
        all_page_translations = [trans_map.get(i, "") for i in range(total_pages)]

        # ── Phase 3: Build ─────────────────────────────────────────
        _update(task_id, step="Building PDF...",
                file_progress=total_pages * 2 - 1, file_total=total_pages * 2)

        _doc = _fitz.open(file_path)
        page_sizes = [(_doc[i].rect.width, _doc[i].rect.height) for i in range(total_pages)]
        _doc.close()

        output_dir = os.path.dirname(file_path)
        output_path = os.path.join(output_dir,
                                   os.path.splitext(os.path.basename(file_path))[0] + "_ocr.pdf")
        output_path = PdfScannedRebuilder._build_clean_text_pdf(
            page_texts=all_page_translations,
            page_sizes=page_sizes,
            output_path=output_path,
            bilingual=bilingual,
            original_pdf=file_path,
        )

        _update(task_id, status="done", step="Complete!",
                output=output_path)

    except Exception as e:
        _update(task_id, status="error", error=f"OCR translation failed: {e}")


def _run_translate_generic(task_id: str, file_path: str, target_lang: str = "zh-CN",
                           bilingual: bool = True, pdf_method: str = "pdf2zh",
                           pages: str = None):
    """Background translation for DOCX/PDF using handlers."""
    try:
        # ── pdf2zh fast path: delegates entirely to PDFMathTranslate ──
        if pdf_method == "pdf2zh" and PdfHandler.is_pdf2zh_available():
            _update(task_id, status="translating", step="Translating with PDFMathTranslate...",
                    current_file=0, total_files=1,
                    file_name=os.path.basename(file_path),
                    file_progress=0, file_total=40)
            try:
                # Run pdf2zh in background thread with heartbeat progress
                import threading
                result_holder = [None]
                error_holder = [None]
                done_flag = threading.Event()
                heartbeat_stop = threading.Event()

                def _run_pdf2zh():
                    try:
                        result_holder[0] = PdfHandler.rebuild_via_pdf2zh(
                            file_path, output_dir=os.path.dirname(file_path),
                            vfont="", vchar="", pages=pages)
                    except Exception as e:
                        error_holder[0] = e
                    done_flag.set()

                def _heartbeat():
                    elapsed = 0
                    while not heartbeat_stop.is_set():
                        heartbeat_stop.wait(1.5)
                        if _tasks.get(task_id, {}).get("stopped"):
                            break  # Stop updating progress
                        elapsed += 1.5
                        # Simulated progress: estimate ~3s per page, cap at 95%
                        pages = _doc_page_count(file_path)
                        est_total = max(pages * 3, 10)
                        pct = min(int(elapsed / est_total * 100), 95)
                        _update(task_id, file_progress=pct, file_total=100,
                                step=f"PDFMathTranslate: {min(pct, 95)}%...")

                t_work = threading.Thread(target=_run_pdf2zh)
                t_beat = threading.Thread(target=_heartbeat)
                t_work.start()
                t_beat.start()
                done_flag.wait()
                heartbeat_stop.set()
                t_beat.join()
                t_work.join()

                if error_holder[0]:
                    _update(task_id, status="error",
                            error=f"PDFMathTranslate failed: {error_holder[0]}")
                    return

                # Check if user stopped during pdf2zh run
                if _tasks.get(task_id, {}).get("stopped"):
                    _update(task_id, status="stopped", step="Stopped by user")
                    return

                output_path = result_holder[0]
            except Exception as e:
                _update(task_id, status="error", error=f"PDFMathTranslate failed: {e}")
                return

            # Move output to output/
            os.makedirs(OUTPUT_DIR, exist_ok=True)
            final_name = os.path.basename(output_path)
            final_path = os.path.join(OUTPUT_DIR, final_name)
            if output_path != final_path:
                shutil.move(output_path, final_path)
                output_path = final_path

            _update(task_id, file_progress=100, file_total=100)
            _update(task_id, status="done",
                    translated=_doc_page_count(output_path) // 2, cached=0,
                    output=output_path, step="Done!")
            return

        _update(task_id, status="loading", step="Loading configuration...")
        config = Config(CONFIG_PATH)
        config.load()
        if not config.api_key or config.api_key in ("sk-xxxx", "sk-your-api-key-here"):
            _update(task_id, status="error", error="API key not configured. Click settings.")
            return

        _update(task_id, step="Extracting document...")
        handler = get_handler(file_path)
        if handler is None:
            _update(task_id, status="error", error="Unsupported file format.")
            return

        fragments = handler.extract(file_path, bilingual=bilingual, pages=pages)

        # ── Scanned PDF → delegate to OCR pipeline ─────────────────────
        if file_path.lower().endswith(".pdf") and config.ocr_enabled:
            from handlers.pdf_ocr import PdfOcr
            if PdfOcr.is_scanned(file_path):
                _run_translate_ocr(task_id, file_path, target_lang, bilingual)
                return

        if not fragments:
            return

        # Source language detection
        sample_text = " ".join(f.text for f in fragments[:5])
        source_lang = detect_language(sample_text)
        _update(task_id, source_lang=source_lang,
                target_lang_name=get_lang_name(target_lang, "zh"))

        # Check if source == target
        from languages.detector import is_same_language
        if is_same_language(sample_text, target_lang):
            _update(task_id, status="error",
                    error=f"Source language ({source_lang}) matches target ({target_lang}). Nothing to translate.")
            return

        # Translate concurrently for speed
        total = len(fragments)
        translator_obj = Translator(config, target_lang, bilingual)

        _update(task_id, status="translating", step="Translating...",
                current_file=0, total_files=1,
                file_name=os.path.basename(file_path),
                file_progress=0, file_total=total)

        texts = [f.text for f in fragments]
        progress_state = {"count": 0}

        # Use cache for DOCX/PDF to support stop & resume
        cache = TranslationCache(os.path.dirname(file_path))
        cache.load()
        uncached_texts, uncached_indices = cache.batch_get(texts)
        total_cached = len(texts) - len(uncached_texts)

        def on_progress(done):
            progress_state["count"] = min(done, total)
            _update(task_id,
                    file_progress=progress_state["count"], file_total=total,
                    step=f"Translating {progress_state['count']}/{total}")

        if uncached_texts:
            try:
                new_translations = translator_obj.translate_all(
                    uncached_texts, progress_callback=on_progress,
                    stop_check=lambda: _tasks.get(task_id, {}).get("stopped", False)
                )
            except StopTranslation:
                _update(task_id, status="stopped", step="Stopped by user")
                return

            # Save new translations to cache
            for text, trans in zip(uncached_texts, new_translations):
                cache.put(text, trans)
            cache.flush()

            # Merge cached + new translations in original order
            all_translations = [None] * len(texts)
            for i, text in enumerate(texts):
                if i not in uncached_indices:
                    all_translations[i] = cache.get(text)
            for idx, trans in zip(uncached_indices, new_translations):
                all_translations[idx] = trans
        else:
            all_translations = [cache.get(t) for t in texts]

        # Retry English-only lines concurrently
        import re
        en_idx = [fi for fi in range(len(all_translations))
                  if not re.search(r'[一-鿿]', str(all_translations[fi]))]
        if en_idx:
            en_texts = [fragments[fi].text for fi in en_idx]
            try:
                retrans = translator_obj.translate_all(en_texts)
                retried = 0
                for fi, trans in zip(en_idx, retrans):
                    if re.search(r'[一-鿿]', str(trans)):
                        all_translations[fi] = trans
                        retried += 1
                if retried:
                    _update(task_id, file_progress=total)
            except Exception:
                pass

        _update(task_id, step="Rebuilding document...")
        output_path = handler.rebuild(
            file_path, fragments, all_translations,
            bilingual=bilingual, target_lang=target_lang,
            method=pdf_method
        )

        _update(task_id, status="done", step="Complete!",
                output=output_path, translated=total, cached=0)

    except Exception as e:
        _update(task_id, status="error", error=str(e))


def _run_translate_babeldoc(task_id: str, file_path: str, target_lang: str = "zh-CN",
                             bilingual: bool = True, pages: str = None):
    """Background PDF translation using BabelDOC engine (one-shot IL pipeline)."""
    try:
        _update(task_id, status="loading", step="Loading configuration...")
        config = Config(CONFIG_PATH)
        config.load()
        if not config.api_key or config.api_key in ("sk-xxxx", "sk-your-api-key-here"):
            _update(task_id, status="error", error="API key not configured. Click settings.")
            return

        _update(task_id, step="Initializing BabelDOC...")

        from handlers.pdf_babeldoc_handler import PdfBabeldocHandler
        handler = PdfBabeldocHandler()

        base_url = config.api_base.rstrip("/")
        if not base_url.endswith("/v1"):
            base_url += "/v1"

        def progress_callback(stage, pct, msg):
            # Don't update progress if user has stopped
            if _tasks.get(task_id, {}).get("stopped"):
                return
            if stage in ("init", "loading"):
                _update(task_id, status="loading", step=msg, file_name="Preparing...")
            elif stage in ("translate", "translating"):
                _update(task_id, status="translating", step=msg,
                        file_name=os.path.basename(file_path),
                        file_progress=pct, file_total=100,
                        total_files=1, current_file=1)
            elif stage == "done":
                _update(task_id, status="done", message=msg)
            elif stage == "error":
                _update(task_id, status="error", error=msg)

        output = handler.translate_full(
            file_path=file_path,
            target_lang=target_lang,
            source_lang="en",
            api_key=config.api_key,
            base_url=base_url,
            model=config.model or "deepseek-chat",
            bilingual=bilingual,
            output_dir=OUTPUT_DIR,
            progress_callback=progress_callback,
            pages=pages,
        )

        # Check if stopped during translation
        if _tasks.get(task_id, {}).get("stopped"):
            _update(task_id, status="stopped", step="Stopped by user")
            return

        _update(task_id, status="done", step="Complete!",
                output=output, translated=0, cached=0)

    except Exception as e:
        _update(task_id, status="error", error=str(e))


# ── API Routes ─────────────────────────────────────────────────────────

@app.post("/api/upload")
async def upload_file(file: UploadFile = File(...)):
    if not is_supported(file.filename):
        exts = ", ".join(get_supported_extensions())
        raise HTTPException(400, f"Unsupported format. Supported: {exts}")

    task_id = uuid.uuid4().hex[:12]
    upload_dir = os.path.join(TEMP_DIR, "_uploads")
    os.makedirs(upload_dir, exist_ok=True)

    safe_name = file.filename
    max_name_len = 80
    if len(safe_name) > max_name_len:
        base, ext = os.path.splitext(safe_name)
        safe_name = base[:max_name_len - len(ext)] + ext
    file_path = os.path.join(upload_dir, f"{task_id}_{safe_name}")

    with open(file_path, "wb") as f:
        shutil.copyfileobj(file.file, f)

    # Read metadata based on file type
    ext = os.path.splitext(file.filename)[1].lower()
    is_scanned = False
    if ext == ".epub":
        try:
            meta = _read_epub_meta(file_path)
        except Exception:
            meta = {"title": file.filename, "author": "", "cover": None, "toc": []}
    elif ext == ".pdf":
        meta = {"title": file.filename, "author": "", "cover": None, "toc": []}
        # Check if it's a scanned PDF (image-based, no text layer)
        try:
            from handlers.pdf_ocr import PdfOcr
            is_scanned = PdfOcr.is_scanned(file_path)
        except Exception:
            pass
    else:
        meta = {"title": file.filename, "author": "", "cover": None, "toc": []}

    _tasks[task_id] = {
        "status": "loaded", "step": "Ready", "task_id": task_id,
        "filename": file.filename, "file_path": file_path,
        "file_type": ext,
        "opf_dir": meta.get("opf_dir", "OEBPS/"),
        "start_time": time.time(),
    }

    return {
        "task_id": task_id,
        "filename": file.filename,
        "file_type": ext,
        "title": meta.get("title", file.filename),
        "author": meta.get("author", ""),
        "cover": meta.get("cover"),
        "toc": meta.get("toc", []),
        "is_scanned": is_scanned,
    }


@app.get("/api/book/{task_id}/content")
async def book_content(task_id: str, path: str = Query(...)):
    task = _tasks.get(task_id)
    if not task:
        raise HTTPException(404, "Task not found")
    epub_path = task.get("file_path")
    if not epub_path or not os.path.exists(epub_path):
        raise HTTPException(404, "EPUB file not found")
    opf_dir = task.get("opf_dir", "OEBPS/")
    html = _read_epub_content(epub_path, opf_dir, path)
    return {"html": html}


@app.get("/api/book/{task_id}/pdf-info")
async def pdf_info(task_id: str):
    """Return PDF metadata: page count, title, file size."""
    task = _tasks.get(task_id)
    if not task:
        raise HTTPException(404, "Task not found")
    file_path = task.get("file_path")
    if not file_path or not os.path.exists(file_path):
        raise HTTPException(404, "File not found")
    ext = task.get("file_type", "")
    if ext != ".pdf":
        raise HTTPException(400, "Not a PDF file")

    import fitz
    try:
        doc = fitz.open(file_path)
        page_count = len(doc)
        meta = doc.metadata or {}
        doc.close()
    except Exception as e:
        raise HTTPException(500, f"Failed to read PDF: {e}")

    return {
        "page_count": page_count,
        "title": meta.get("title", task.get("filename", "")),
        "file_size": os.path.getsize(file_path),
    }


@app.get("/api/book/{task_id}/pdf-page")
async def pdf_page(task_id: str, page: int = Query(..., ge=0), scale: float = Query(1.5, ge=0.5, le=4.0)):
    """Render a single PDF page as a PNG image."""
    task = _tasks.get(task_id)
    if not task:
        raise HTTPException(404, "Task not found")
    file_path = task.get("file_path")
    if not file_path or not os.path.exists(file_path):
        raise HTTPException(404, "File not found")

    import fitz
    try:
        doc = fitz.open(file_path)
        if page < 0 or page >= len(doc):
            doc.close()
            raise HTTPException(404, f"Page {page} out of range (0-{len(doc)-1})")
        p = doc[page]
        mat = fitz.Matrix(scale, scale)
        pix = p.get_pixmap(matrix=mat)
        img_bytes = pix.tobytes("png")
        doc.close()
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(500, f"Failed to render page: {e}")

    return StreamingResponse(io.BytesIO(img_bytes), media_type="image/png",
                             headers={"Cache-Control": "public, max-age=3600"})


@app.get("/api/book/{task_id}/pdf-text")
async def pdf_text(task_id: str, page: int = Query(..., ge=0)):
    """Extract text from a single PDF page."""
    task = _tasks.get(task_id)
    if not task:
        raise HTTPException(404, "Task not found")
    file_path = task.get("file_path")
    if not file_path or not os.path.exists(file_path):
        raise HTTPException(404, "File not found")

    import fitz
    try:
        doc = fitz.open(file_path)
        if page < 0 or page >= len(doc):
            doc.close()
            raise HTTPException(404, f"Page {page} out of range (0-{len(doc)-1})")
        p = doc[page]
        text = p.get_text()
        doc.close()
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(500, f"Failed to extract text: {e}")

    return {"page": page, "text": text}


@app.get("/api/book/{task_id}/docx-content")
async def docx_content(task_id: str):
    """Read DOCX content and return as simple HTML for preview."""
    task = _tasks.get(task_id)
    if not task:
        raise HTTPException(404, "Task not found")
    file_path = task.get("file_path")
    if not file_path or not os.path.exists(file_path):
        raise HTTPException(404, "File not found")
    if task.get("file_type") != ".docx":
        raise HTTPException(400, "Not a DOCX file")

    try:
        import docx
        doc = docx.Document(file_path)
    except Exception as e:
        raise HTTPException(500, f"Failed to read DOCX: {e}")

    html_parts = ['<div class="docx-preview">']

    # Collect paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            html_parts.append('<p style="margin-bottom:0.5em;">&nbsp;</p>')
            continue
        style_name = para.style.name if para.style else "Normal"

        # Determine tag and styling
        if style_name.startswith("Heading 1"):
            tag = "h1"
        elif style_name.startswith("Heading 2"):
            tag = "h2"
        elif style_name.startswith("Heading 3"):
            tag = "h3"
        elif style_name.startswith("Heading"):
            tag = "h4"
        else:
            tag = "p"

        # Build HTML with run-level formatting
        inner = ""
        for run in para.runs:
            t = run.text
            if not t:
                continue
            if run.bold:
                t = f"<strong>{t}</strong>"
            if run.italic:
                t = f"<em>{t}</em>"
            if run.underline:
                t = f"<u>{t}</u>"
            inner += t

        if not inner:
            inner = text

        html_parts.append(f"<{tag} style='margin-bottom:0.5em;'>{inner}</{tag}>")

    # Collect tables
    for ti, table in enumerate(doc.tables):
        html_parts.append(f'<table style="border-collapse:collapse;width:100%;margin:1em 0;">')
        for ri, row in enumerate(table.rows):
            html_parts.append("<tr>")
            for ci, cell in enumerate(row.cells):
                tag = "th" if ri == 0 else "td"
                html_parts.append(f'<{tag} style="border:1px solid #ccc;padding:6px 10px;text-align:left;vertical-align:top;">{cell.text}</{tag}>')
            html_parts.append("</tr>")
        html_parts.append("</table>")

    html_parts.append("</div>")
    return {"html": "\n".join(html_parts), "title": task.get("filename", "")}


@app.post("/api/start/{task_id}")
async def start_translation(task_id: str, body: dict = None):
    task = _tasks.get(task_id)
    if not task:
        raise HTTPException(404, "Task not found")
    if task.get("status") == "translating":
        raise HTTPException(400, "Already translating")
    file_path = task.get("file_path") or task.get("epub_path")
    if not file_path or not os.path.exists(file_path):
        raise HTTPException(404, "File not found")

    # Get target language and mode from request body
    target_lang = "zh-CN"
    bilingual = True
    pdf_method = "babeldoc"  # default PDF engine
    pages = None  # page range string, e.g. "1-5,8"
    if body:
        target_lang = body.get("target_lang", target_lang)
        bilingual = body.get("bilingual", True)
        pdf_method = body.get("pdf_method", pdf_method)
        pages = body.get("pages") or None

    _update(task_id, status="queued", step="Starting...", stopped=False,
            target_lang=target_lang, bilingual=bilingual)

    file_type = task.get("file_type", ".epub")
    if file_type == ".epub":
        _executor.submit(_run_translate, task_id, file_path, target_lang, bilingual)
    else:
        _executor.submit(_run_translate_generic, task_id, file_path,
                         target_lang, bilingual, pdf_method, pages)

    return {"ok": True}


@app.post("/api/start_babeldoc/{task_id}")
async def start_babeldoc_translation(task_id: str, body: dict = None):
    task = _tasks.get(task_id)
    if not task:
        raise HTTPException(404, "Task not found")
    if task.get("status") == "translating":
        raise HTTPException(400, "Already translating")
    file_path = task.get("file_path")
    if not file_path or not os.path.exists(file_path):
        raise HTTPException(404, "File not found")

    target_lang = "zh-CN"
    bilingual = True
    pages = None
    if body:
        target_lang = body.get("target_lang", target_lang)
        bilingual = body.get("bilingual", True)
        pages = body.get("pages") or None

    # ── Scanned PDF detection ──────────────────────────────────────
    if file_path.lower().endswith(".pdf"):
        try:
            from handlers.pdf_ocr import PdfOcr
            config = Config(CONFIG_PATH)
            config.load()
            if config.ocr_enabled and PdfOcr.is_scanned(file_path):
                ui_lang = (body or {}).get("ui_lang", "zh")
                if ui_lang.startswith("zh"):
                    msg = ("此 PDF 疑似扫描版（图片型）文档，BabelDOC 仅支持文字型 PDF。"
                           "请使用 OCR 翻译页面（/ocrtranslate）处理扫描版 PDF。")
                else:
                    msg = ("This PDF appears to be a scanned (image-based) document. "
                           "BabelDOC only supports text-based PDFs. "
                           "Please use the OCR Translation page (/ocrtranslate) for scanned PDFs.")
                raise HTTPException(400, msg)
        except HTTPException:
            raise
        except Exception:
            pass  # If scan detection fails, proceed with BabelDOC

    _update(task_id, status="queued", step="Starting BabelDOC...", stopped=False,
            target_lang=target_lang, bilingual=bilingual)

    _executor.submit(_run_translate_babeldoc, task_id, file_path,
                     target_lang, bilingual, pages)

    return {"ok": True}


@app.post("/api/start_ocr/{task_id}")
async def start_ocr_translation(task_id: str, body: dict = None):
    task = _tasks.get(task_id)
    if not task:
        raise HTTPException(404, "Task not found")
    if task.get("status") == "translating":
        raise HTTPException(400, "Already translating")
    file_path = task.get("file_path")
    if not file_path or not os.path.exists(file_path):
        raise HTTPException(404, "File not found")

    target_lang = "zh-CN"
    bilingual = True
    if body:
        target_lang = body.get("target_lang", target_lang)
        bilingual = body.get("bilingual", True)

    _update(task_id, status="queued", step="Starting OCR translation...", stopped=False,
            target_lang=target_lang, bilingual=bilingual)

    _executor.submit(_run_translate_ocr, task_id, file_path,
                     target_lang, bilingual)

    return {"ok": True}


@app.post("/api/stop/{task_id}")
async def stop_translation(task_id: str):
    task = _tasks.get(task_id)
    if not task:
        raise HTTPException(404, "Task not found")
    if task.get("status") not in ("translating", "queued", "loading"):
        raise HTTPException(400, "Not running")
    _update(task_id, stopped=True)
    return {"ok": True}


@app.get("/api/task/{task_id}")
async def get_task_status(task_id: str):
    """Return current task state as JSON (for resume after page navigation)."""
    task = _tasks.get(task_id)
    if not task:
        raise HTTPException(404, "Task not found")
    return dict(task)


@app.get("/api/progress/{task_id}")
async def progress_stream(task_id: str):
    if task_id not in _tasks:
        raise HTTPException(404, "Task not found")

    async def event_gen():
        last_data = None
        while True:
            with _lock:
                task = dict(_tasks.get(task_id, {}))
            data = json.dumps(task, ensure_ascii=False)
            if data != last_data:
                yield f"data: {data}\n\n"
                last_data = data
            if task.get("status") in ("done", "error", "stopped"):
                break
            await asyncio.sleep(0.4)

    return StreamingResponse(event_gen(), media_type="text/event-stream",
                             headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"})


@app.get("/api/download/{task_id}")
async def download(task_id: str):
    task = _tasks.get(task_id)
    if not task or task.get("status") != "done":
        raise HTTPException(404, "File not ready")
    output = task.get("output")
    if not output or not os.path.exists(output):
        raise HTTPException(404, "Output file not found")
    return FileResponse(output, filename=os.path.basename(output), media_type="application/epub+zip")


# ── Language API ────────────────────────────────────────────────────

@app.get("/api/languages")
async def list_languages(ui: str = Query(default="zh")):
    return get_all_languages(ui)


@app.get("/api/formats")
async def list_formats():
    return {"extensions": get_supported_extensions()}


# ── Config API ─────────────────────────────────────────────────────────

def _mask_key(key: str) -> str:
    """Mask API key for display. Handles both plain and encrypted keys."""
    if is_encrypted(key):
        key = decrypt(key)
    if not key or len(key) < 8:
        return ""
    return key[:6] + "****" + key[-4:]


@app.get("/api/config")
async def get_config():
    try:
        with open(CONFIG_PATH, "r", encoding="utf-8") as f:
            cfg = yaml.safe_load(f) or {}
    except FileNotFoundError:
        cfg = {}

    raw_key = cfg.get("api_key", "")
    encrypted = is_encrypted(raw_key)
    # Decrypt if needed, then check if it's a real key (not placeholder)
    plain_key = decrypt(raw_key) if encrypted else raw_key
    has_key = bool(plain_key and plain_key not in ("sk-xxxx", "sk-your-api-key-here"))

    # OCR API key
    raw_ocr_key = cfg.get("ocr_api_key", "")
    ocr_encrypted = is_encrypted(raw_ocr_key)
    plain_ocr_key = decrypt(raw_ocr_key) if ocr_encrypted else raw_ocr_key
    has_ocr_key = bool(plain_ocr_key)

    return {
        "api_key_masked": _mask_key(raw_key) if has_key else "",
        "api_key_set": has_key,
        "api_key_encrypted": encrypted,
        "api_base": cfg.get("api_base", ""),
        "model": cfg.get("model", ""),
        "translation_mode": cfg.get("translation_mode", "bilingual"),
        "ocr_enabled": cfg.get("ocr_enabled", True),
        "ocr_api_key_masked": _mask_key(raw_ocr_key) if has_ocr_key else "",
        "ocr_api_key_set": has_ocr_key,
        "ocr_api_key_encrypted": ocr_encrypted,
        "ocr_api_base": cfg.get("ocr_api_base", ""),
        "ocr_model": cfg.get("ocr_model", ""),
        "max_file_size_mb": cfg.get("max_file_size_mb", 500),
        "max_concurrent_tasks": cfg.get("max_concurrent_tasks", 1),
    }


@app.post("/api/config")
async def update_config(body: dict):
    try:
        with open(CONFIG_PATH, "r", encoding="utf-8") as f:
            cfg = yaml.safe_load(f) or {}
    except FileNotFoundError:
        cfg = {}

    if "api_key" in body:
        raw_key = body["api_key"].strip()
        # Encrypt the key before storing
        cfg["api_key"] = encrypt(raw_key) if raw_key else ""
    if "api_base" in body:
        cfg["api_base"] = body["api_base"].strip()
    if "model" in body:
        cfg["model"] = body["model"].strip()
    if "translation_mode" in body:
        cfg["translation_mode"] = body["translation_mode"]
    # OCR settings
    if "ocr_enabled" in body:
        cfg["ocr_enabled"] = body["ocr_enabled"]
    if "ocr_api_key" in body:
        raw_ocr = body["ocr_api_key"].strip()
        cfg["ocr_api_key"] = encrypt(raw_ocr) if raw_ocr else ""
    if "ocr_api_base" in body:
        cfg["ocr_api_base"] = body["ocr_api_base"].strip()
    if "ocr_model" in body:
        cfg["ocr_model"] = body["ocr_model"].strip()
    # General settings
    if "max_file_size_mb" in body:
        cfg["max_file_size_mb"] = int(body["max_file_size_mb"])
    if "max_concurrent_tasks" in body:
        cfg["max_concurrent_tasks"] = int(body["max_concurrent_tasks"])

    with open(CONFIG_PATH, "w", encoding="utf-8") as f:
        yaml.dump(cfg, f, default_flow_style=False, allow_unicode=True)

    return {"ok": True, "api_key_masked": _mask_key(cfg.get("api_key", ""))}


# ── Frontend ───────────────────────────────────────────────────────────

from fastapi.staticfiles import StaticFiles

_images_dir = os.path.join(os.path.dirname(__file__), "images")
if os.path.isdir(_images_dir):
    app.mount("/images", StaticFiles(directory=_images_dir), name="images")


@app.get("/favicon.svg")
async def favicon_svg():
    favicon_path = os.path.join(os.path.dirname(__file__), "favicon.svg")
    with open(favicon_path, "r", encoding="utf-8") as f:
        return Response(content=f.read(), media_type="image/svg+xml")


@app.get("/favicon.png")
async def favicon_png():
    favicon_path = os.path.join(os.path.dirname(__file__), "favicon.png")
    return FileResponse(favicon_path, media_type="image/png")


@app.get("/", response_class=HTMLResponse)
async def index():
    html_path = os.path.join(os.path.dirname(__file__), "index.html")
    with open(html_path, "r", encoding="utf-8") as f:
        return f.read()


@app.get("/guide", response_class=HTMLResponse)
async def guide():
    guide_path = os.path.join(os.path.dirname(__file__), "guide.html")
    with open(guide_path, "r", encoding="utf-8") as f:
        return f.read()


@app.get("/translate", response_class=HTMLResponse)
async def translate_page():
    path = os.path.join(os.path.dirname(__file__), "translate.html")
    with open(path, "r", encoding="utf-8") as f:
        return f.read()


@app.get("/ocrtranslate", response_class=HTMLResponse)
async def ocr_translate_page():
    path = os.path.join(os.path.dirname(__file__), "ocrtranslate.html")
    with open(path, "r", encoding="utf-8") as f:
        return f.read()


@app.get("/login", response_class=HTMLResponse)
async def login_page():
    path = os.path.join(os.path.dirname(__file__), "login.html")
    with open(path, "r", encoding="utf-8") as f:
        return f.read()


@app.get("/batch", response_class=HTMLResponse)
async def batch_page():
    path = os.path.join(os.path.dirname(__file__), "batch.html")
    with open(path, "r", encoding="utf-8") as f:
        return f.read()


@app.get("/settings-modal.js")
async def settings_modal_js():
    """Serve the shared settings modal JavaScript."""
    js_path = os.path.join(os.path.dirname(__file__), "settings-modal.js")
    with open(js_path, "r", encoding="utf-8") as f:
        return Response(content=f.read(), media_type="application/javascript")
